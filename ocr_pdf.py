from mistralai import Mistral
import os
import json
import base64
from docx import Document
from docx.shared import Inches
import re

# Set API key
api_key = os.environ["MISTRAL_API_KEY"]

# Initialize Mistral client
client = Mistral(api_key=api_key)

# Upload PDF file for OCR
file_path = r"C:\Users\thamb\Downloads\Machine Learning by TutorialsPoint.pdf"
uploaded_pdf = client.files.upload(
    file={
        "file_name": "sample",
        "content": open(file_path, "rb"),
    },
    purpose="ocr"
)

print("Uploaded File:", uploaded_pdf)

# Get signed URL
signed_url = client.files.get_signed_url(file_id=uploaded_pdf.id)

# Process OCR
ocr_response = client.ocr.process(
    model="mistral-ocr-latest",
    document={
        "type": "document_url",
        "document_url": signed_url.url,
    },
    include_image_base64=True
)

# Function to serialize OCR response
def serialize(obj):
    """Recursively convert objects to dictionaries for JSON serialization."""
    if isinstance(obj, list):
        return [serialize(item) for item in obj]
    elif hasattr(obj, "__dict__"):
        return {key: serialize(value) for key, value in obj.__dict__.items()}
    else:
        return obj

# Extract and serialize OCR pages
pages = ocr_response.pages
pages_data = serialize(pages)

# Save OCR response as JSON
json_output_file = "ocr_output.json"
with open(json_output_file, "w", encoding="utf-8") as json_file:
    json.dump(pages_data, json_file, indent=4, ensure_ascii=False)

print(f"OCR response saved as {json_output_file}")



def json_to_doc(json_data, output_file="ocr_output.docx"):
    doc = Document()
    
    for page in json_data:
        doc.add_paragraph(page.get("markdown", ""))  # Extract text
        
        for image in page.get("images", []):
            image_base64 = image.get("image_base64", "")
            image_id = image.get("id", "Unknown")
            
            if image_base64:
                try:
                    print(f"Processing Image ID: {image_id}")
                    
                    # Remove the "data:image/jpeg;base64," prefix
                    image_base64 = re.sub(r"^data:image\/\w+;base64,", "", image_base64)
                    
                    # Decode base64 image
                    image_data = base64.b64decode(image_base64)
                    
                    # Save image temporarily
                    image_file = f"temp_{image_id}.jpg"
                    with open(image_file, "wb") as f:
                        f.write(image_data)
                    
                    # Add image to document
                    doc.add_picture(image_file, width=Inches(5))
                    
                    # Add image name below the image
                    doc.add_paragraph(f"Image: {image_file}")
                
                except Exception as e:
                    print(f"Error processing image: {e}")
    
    # Save the document
    doc.save(output_file)
    print(f"Document saved as {output_file}")




json_to_doc(pages_data, "ocr_output.docx")
